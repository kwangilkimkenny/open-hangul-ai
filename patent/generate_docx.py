#!/usr/bin/env python3
"""Generate Word document from patent markdown"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PATENT_DIR = os.path.dirname(__file__)
FIGURES_DIR = os.path.join(PATENT_DIR, 'figures')

doc = Document()

# ============================================================
# Style setup
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'AppleGothic'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

# Set CJK font
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), 'AppleGothic')

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'AppleGothic'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), 'AppleGothic')
    return h


def add_para(text, bold=False, italic=False, size=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'AppleGothic'
    rPr = run._element.get_or_add_rPr()
    rFonts_el = rPr.find(qn('w:rFonts'))
    if rFonts_el is None:
        from docx.oxml import OxmlElement
        rFonts_el = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts_el)
    rFonts_el.set(qn('w:eastAsia'), 'AppleGothic')
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(filename, caption):
    """Add figure image with caption"""
    path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.italic = True
        cap.paragraph_format.space_after = Pt(12)


def add_table(headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'AppleGothic'

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'AppleGothic'

    doc.add_paragraph()  # spacer


# ============================================================
# Document Content
# ============================================================

# Title page
doc.add_paragraph()
doc.add_paragraph()
add_para('특 허 명 세 서', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
doc.add_paragraph()
add_para('AI 기반 보안 검증 및 규제 준수 기능이 통합된\n지능형 문서 편집 시스템 및 그 방법', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
doc.add_paragraph()
add_para('Intelligent Document Editing System and Method with Integrated\nAI-Based Security Verification and Regulatory Compliance', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
doc.add_paragraph()
doc.add_paragraph()
add_para('출원인: [출원인 정보]', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('발명자: [발명자 정보]', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('출원일: 2026년    월    일', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 기술분야
# ============================================================
add_heading('[기술분야]', level=1)
add_para('본 발명은 인공지능(AI) 기반의 문서 편집 기술에 관한 것으로, 보다 구체적으로는 대규모 언어 모델(LLM)을 활용한 콘텐츠 자동 생성, 다계층 보안 방어, 할루시네이션(환각) 검증, 및 다중 규제 프레임워크 준수 평가가 단일 문서 편집 환경 내에서 유기적으로 통합된 지능형 문서 편집 시스템 및 그 운용 방법에 관한 것이다.')
add_para('본 발명은 국제특허분류(IPC) G06F 40/00(디지털 데이터 처리: 자연어 데이터 처리), G06F 21/00(보안 장치), G06N 3/08(기계 학습) 분야에 속한다.')

# ============================================================
# 배경기술
# ============================================================
add_heading('[배경기술]', level=1)
add_heading('1. 종래 기술의 현황', level=2)
add_para('최근 대규모 언어 모델(Large Language Model, 이하 "LLM")의 급속한 발전에 따라, 문서 편집 분야에서도 AI를 활용한 자동 콘텐츠 생성 기능이 도입되고 있다. 대표적으로 마이크로소프트 코파일럿(Microsoft Copilot), 구글 제미니(Google Gemini) 등이 기존 오피스 스위트에 AI 기능을 탑재하여 문서 초안 작성, 요약, 번역 등의 기능을 제공하고 있다.')
add_para('그러나 종래의 AI 통합 문서 편집 시스템은 다음과 같은 기술적 한계를 가지고 있다.')

add_heading('1.1 보안 취약성', level=3)
add_para('종래 시스템은 LLM에 대한 프롬프트 인젝션(Prompt Injection), 탈옥 공격(Jailbreak Attack), 사회공학적 공격(Social Engineering) 등 AI 특유의 보안 위협에 대한 체계적인 방어 메커니즘이 부재하다. 특히, 한국어, 일본어, 중국어 등 CJK 언어권에서의 언어학적 우회 공격(자모 분리, 초성 약어, 고어 활용, 동형 문자 치환 등)에 대한 방어 기술은 전무한 실정이다.')

add_heading('1.2 할루시네이션 문제', level=3)
add_para('LLM이 생성한 콘텐츠에는 사실과 다른 정보가 포함될 수 있는 할루시네이션(Hallucination) 현상이 불가피하게 발생한다. 종래 시스템은 생성된 콘텐츠의 사실적 정확성을 원본 문서와 대조하여 실시간으로 검증하는 메커니즘이 없거나, 있더라도 단일 계층의 단순 비교에 그치고 있어 수치 오류, 문맥 왜곡, 규제 위반 표현 등을 포괄적으로 탐지하지 못한다.')

add_heading('1.3 규제 준수 부재', level=3)
add_para('2024년 EU AI Act(인공지능법) 발효, 2026년 한국 인공지능 기본법 시행, NIST AI RMF(위험관리 프레임워크) 공포 등 AI 관련 국제 규제가 급속히 강화되고 있다. 그러나 종래의 문서 편집 시스템은 AI 사용 이력의 체계적 기록, 투명성 보장, 위험 평가, 인간 감독(Human Oversight) 등 규제 요건을 충족하기 위한 기술적 수단을 내장하고 있지 않다.')

add_heading('1.4 기능 분절 문제', level=3)
add_para('종래 기술에서 보안, 팩트체크, 규제 준수는 각각 독립된 별개의 도구로 제공되어, 문서 편집 워크플로우와 유기적으로 연동되지 않는다. 이로 인해 사용자는 복수의 도구를 번갈아 사용해야 하며, 데이터 흐름의 단절로 인한 정보 손실과 업무 효율 저하가 발생한다.')

add_heading('2. 선행기술 문헌', level=2)
add_para('• 특허문헌 1: US 2023/0186001 A1 — "AI-Assisted Document Generation System" (Microsoft, 2023)')
add_para('• 특허문헌 2: US 2024/0028742 A1 — "Methods for Detecting AI-Generated Content" (Google, 2024)')
add_para('• 비특허문헌 1: NVIDIA NeMo Guardrails (2023) — LLM 가드레일 오픈소스')
add_para('• 비특허문헌 2: Guardrails AI (2023) — LLM 출력 검증 프레임워크')

# ============================================================
# 발명이 해결하고자 하는 과제
# ============================================================
add_heading('[발명이 해결하고자 하는 과제]', level=1)
add_para('본 발명은 상기한 종래 기술의 문제점을 해결하기 위하여 안출된 것으로, 다음의 기술적 과제를 해결하는 것을 목적으로 한다.')
add_para('제1 과제: 문서 편집 환경에서 LLM 기반 콘텐츠 생성 시, 한국어를 포함한 CJK 언어 특화 공격을 방어할 수 있는 다계층 보안 방어 시스템을 제공한다.', bold=True)
add_para('제2 과제: AI가 생성한 콘텐츠의 사실적 정확성을 원본 문서와 대조하여 규칙 기반, 수치 검증, 자연어 추론(NLI), LLM 재검증의 4계층으로 실시간 검증하는 할루시네이션 탐지 메커니즘을 제공한다.', bold=True)
add_para('제3 과제: EU AI Act, 한국 AI 기본법, NIST AI RMF, OWASP LLM Top 10 등 다중 규제 프레임워크에 대한 준수 여부를 AI 활동 로그 기반으로 자동 평가하고 증거 기반 리포트를 생성하는 컴플라이언스 시스템을 제공한다.', bold=True)
add_para('제4 과제: 상기 보안 방어, 할루시네이션 검증, 규제 준수 평가를 문서 편집 워크플로우 내에 유기적으로 통합하여, 단일 사용자 인터페이스에서 전체 파이프라인을 운용할 수 있는 통합 시스템을 제공한다.', bold=True)
add_para('제5 과제: 서버 장애 또는 오프라인 환경에서도 핵심 보안 및 검증 기능이 유지되는 온라인-오프라인 하이브리드 아키텍처를 제공한다.', bold=True)

# ============================================================
# 과제의 해결 수단
# ============================================================
add_heading('[과제의 해결 수단]', level=1)
add_para('상기 과제를 해결하기 위하여 본 발명은 다음의 기술적 수단을 제공한다.')
add_para('본 발명의 일 실시예에 따른 AI 기반 보안 검증 및 규제 준수 기능이 통합된 지능형 문서 편집 시스템은:', bold=True)
add_para('(a) 문서 파싱 및 편집 모듈 — 한글 문서 포맷(HWPX/HWP)을 포함하는 문서 파일을 파싱하여 내부 문서 객체 모델로 변환하고, 사용자의 편집 입력을 수신하여 상기 문서 객체 모델을 갱신하는 모듈;')
add_para('(b) AI 콘텐츠 생성 모듈 — 상기 문서 객체 모델의 구조를 분석하여 문서 유형을 자동 판별하고, 헤더-콘텐츠 쌍 및 테이블 관계를 추출하여 문맥 인식 프롬프트를 구성한 후, LLM API를 호출하여 콘텐츠를 생성하고 병합하는 모듈;')
add_para('(c) 다계층 보안 방어 모듈(AEGIS) — 6개 방어 계층을 순차 적용하며, 한국어 언어학적 우회 공격 특화 12개 탐지 서브모듈을 포함하는 모듈;')
add_para('(d) 할루시네이션 검증 모듈(TruthAnchor) — 4계층 검증 파이프라인과 5차원 신뢰도 스코어링을 수행하는 모듈;')
add_para('(e) 규제 준수 평가 모듈 — 4대 규제 프레임워크에 대한 자동 평가 및 증거 기반 리포트를 생성하는 모듈;')
add_para('(f) 통합 제어 모듈 — 전체 파이프라인을 자동 오케스트레이션하는 모듈;')
add_para('을 포함하는 것을 특징으로 한다.')

# ============================================================
# 발명의 효과
# ============================================================
add_heading('[발명의 효과]', level=1)
add_para('첫째, 6계층 PALADIN 방어 아키텍처와 한국어 특화 12개 탐지 모듈을 통해, 종래 시스템에서 방어가 불가능했던 CJK 언어학적 우회 공격을 효과적으로 탐지 및 차단할 수 있다.', bold=True)
add_para('둘째, 4계층 할루시네이션 검증 파이프라인과 5차원 신뢰도 스코어링을 통해, AI 생성 콘텐츠의 사실적 정확성을 다각적으로 검증할 수 있다.', bold=True)
add_para('셋째, 4대 국제 규제 프레임워크에 대한 자동화된 준수 평가 및 증거 기반 리포트 생성을 통해, 규제 감사 대응 비용을 대폭 절감할 수 있다.', bold=True)
add_para('넷째, 보안-검증-규제 준수가 문서 편집 워크플로우에 통합되어 업무 비효율을 해소한다.', bold=True)
add_para('다섯째, 온라인-오프라인 하이브리드 아키텍처를 통해 핵심 안전 기능의 연속성이 보장된다.', bold=True)

# ============================================================
# 도면의 간단한 설명
# ============================================================
add_heading('[도면의 간단한 설명]', level=1)

figures_desc = [
    ('도 1', '본 발명의 일 실시예에 따른 지능형 문서 편집 시스템의 전체 아키텍처 블록도'),
    ('도 2', '본 발명의 일 실시예에 따른 AI 콘텐츠 생성 및 검증 파이프라인의 데이터 흐름도'),
    ('도 3', '본 발명의 일 실시예에 따른 다계층 보안 방어 모듈(AEGIS)의 6계층 PALADIN 방어 아키텍처 상세도'),
    ('도 4', '본 발명의 일 실시예에 따른 한국어 언어학적 우회 공격 탐지 서브모듈의 구성도'),
    ('도 5', '본 발명의 일 실시예에 따른 할루시네이션 검증 모듈(TruthAnchor)의 4계층 검증 파이프라인 상세도'),
    ('도 6', '본 발명의 일 실시예에 따른 5차원 할루시네이션 신뢰도 스코어링 모델 구성도'),
    ('도 7', '본 발명의 일 실시예에 따른 다중 규제 프레임워크 컴플라이언스 평가 시스템의 구성도'),
    ('도 8', '본 발명의 일 실시예에 따른 온라인-오프라인 하이브리드 운용 모드 전환 흐름도'),
    ('도 9', '본 발명의 일 실시예에 따른 통합 제어 모듈의 오케스트레이션 시퀀스 다이어그램'),
    ('도 10', '본 발명의 일 실시예에 따른 문서 구조 분석 기반 문맥 인식 프롬프트 구성 방법의 흐름도'),
]

for fig_num, fig_desc in figures_desc:
    add_para(f'[{fig_num}] {fig_desc}')

# ============================================================
# 발명을 실시하기 위한 구체적인 내용
# ============================================================
doc.add_page_break()
add_heading('[발명을 실시하기 위한 구체적인 내용]', level=1)
add_para('이하, 첨부된 도면을 참조하여 본 발명의 바람직한 실시예를 상세히 설명한다.')

# Figure 1
add_heading('1. 시스템 전체 구성 (도 1 참조)', level=2)
add_para('본 발명의 일 실시예에 따른 지능형 문서 편집 시스템(100)은 도 1에 도시된 바와 같이, 문서 파싱 및 편집 모듈(110), AI 콘텐츠 생성 모듈(120), 다계층 보안 방어 모듈(130), 할루시네이션 검증 모듈(140), 규제 준수 평가 모듈(150), 및 통합 제어 모듈(160)을 포함하여 구성된다.')
add_figure('fig01_system_architecture.png', '[도 1] 지능형 문서 편집 시스템 전체 아키텍처 블록도')

add_heading('1.1 문서 파싱 및 편집 모듈 (110)', level=3)
add_para('문서 파서(111)는 HWPX 포맷의 압축 아카이브를 해제하고, 내부 XML 구조(content.xml, docinfo.xml)를 파싱하여 계층적 문서 객체 모델로 변환한다. 상기 문서 객체 모델은 섹션, 단락, 테이블, 이미지, 도형 등의 요소를 트리 구조로 표현하며, 각 요소는 서식 속성을 포함한다.')
add_para('인라인 편집기(113)는 셀 단위 및 단락 단위의 실시간 편집을 지원하며, 이력 관리부(114)는 스택 기반의 실행 취소/재실행 메커니즘을 제공한다.')

add_heading('1.2 AI 콘텐츠 생성 모듈 (120)', level=3)
add_para('문서 구조 분석부(121)는 문서 객체 모델을 입력받아 문서 유형을 자동 판별하고, 헤더-콘텐츠 쌍을 추출하며, 테이블 내 열/행 간의 의미적 관계를 분석한다. 프롬프트 빌더(122)는 분석 결과를 기반으로 문맥 인식 프롬프트를 구성한다.')

# Figure 2
add_figure('fig02_pipeline_dataflow.png', '[도 2] AI 콘텐츠 생성 및 검증 파이프라인 데이터 흐름도')

# AEGIS
add_heading('1.3 다계층 보안 방어 모듈 — AEGIS (130)', level=2)
add_para('다계층 보안 방어 모듈(130)은 6계층 PALADIN 방어 오케스트레이터를 핵심으로 하며, 각 계층은 순차적으로 실행되어 누적적 위험 점수를 산출한다.')

add_table(
    ['계층', '명칭', '기능'],
    [
        ['제0계층', '신뢰 경계 검증부 (131)', '유니코드 정규화, 제로폭 문자 탐지, 동형 문자 정규화'],
        ['제1계층', '의도 분석부 (132)', '탈옥, 프롬프트 인젝션, 사회공학 공격 패턴 분석'],
        ['제2계층', 'RA 가드 (133)', 'RAG 환경 간접 인젝션, 악성 컨텍스트 탐지'],
        ['제3계층', 'ClassRAG (134)', '문맥 분류 기반 컨텍스트 일관성 검증'],
        ['제4계층', '회로 차단기 (135)', '비율 제한, 누적 위험 점수 기반 자동 차단'],
        ['제5계층', '행동 분석부 (136)', '크레센도 공격 탐지, 세션 기반 위험도 누적'],
    ]
)

add_figure('fig03_paladin_architecture.png', '[도 3] 6계층 PALADIN 방어 아키텍처 상세도')

add_para('행동 분석부(136)는 다회차 대화에서의 크레센도 공격을 탐지한다: 패턴 A(일반→전문→구체), 패턴 B(교육→심화→유해), 패턴 C(신뢰→역할→제약해제). 각 계층은 결정론적 방어 관리자(θ=0 보장)를 포함한다.')

# Korean modules
add_heading('1.3.1 한국어 언어학적 공격 탐지 서브모듈 (137)', level=3)
add_table(
    ['번호', '탐지기', '탐지 대상'],
    [
        ['137-1', '키보드 매핑 탐지기', '한국어 자판 배열 인코딩 공격'],
        ['137-2', '고어 탐지기', '고대 한국어 문자 필터 우회'],
        ['137-3', '동형 문자 탐지기', '시각적 동일 한국어 문자 치환'],
        ['137-4', '한자 탐지기', '한자 주입 우회'],
        ['137-5', '토크나이저 취약점 탐지기', '토큰 경계 악용'],
        ['137-6', '자모 정규화기', '한글 자모 분해 우회'],
        ['137-7', '은어 탐지기', '은어/속어 유해 표현'],
        ['137-8', '조사 탐지기', '형태소 조사 의미 변조'],
        ['137-9', '음절 역순 탐지기', '문자 순서 역전 우회'],
        ['137-10', '음운 변이 탐지기', '유사 발음 대체 우회'],
        ['137-11', '초성 디코더', '초성 약어 (ㅈㅅ→죄송) 유해 표현'],
        ['137-12', '코드 스위칭 탐지기', '혼합 언어 공격'],
    ]
)

add_figure('fig04_korean_submodules.png', '[도 4] 한국어 언어학적 우회 공격 탐지 서브모듈 구성도')

# Risk score
add_para('위험 점수 산출: rawScore = min(layerScore + koreanScore + piiScore, 100), finalScore = min(rawScore × sensitivity, 100). 위험 수준은 critical(60+), high(40~60), medium(20~40), low(<20)로 분류된다.', bold=True)

# TruthAnchor
doc.add_page_break()
add_heading('1.4 할루시네이션 검증 모듈 — TruthAnchor (140)', level=2)
add_para('주장 추출부(141)는 AI 생성 텍스트에서 검증 가능한 개별 주장을 추출한다. 한국어 특화 문장 분리 및 복합문 분할을 수행하며, 비사실적 표현은 필터링한다.')
add_para('증거 매칭부(142)는 인메모리 문자열 유사도 기반 매칭을 수행한다: 키워드 겹침(50%) + 문자 n-그램(30%) + 시퀀스 유사도(20%).')

add_table(
    ['검증 계층', '명칭', '설명', '지연시간'],
    [
        ['제0계층', '가드레일 검증부', '38개 도메인 규칙 (규제 위반 탐지)', '<1ms'],
        ['제0.5계층', '수치 교차검증부', '문맥 카테고리 기반 수치 비교', '<1ms'],
        ['제1계층', 'NLI 검증부', 'DeBERTa-v3 cross-encoder (의미 검증)', '~50ms'],
        ['제2계층', 'LLM 재검증부', '중립 판정 주장에 대한 GPT-4 재검증', '~2s'],
    ]
)

add_figure('fig05_truthanchor_pipeline.png', '[도 5] TruthAnchor 4계층 검증 파이프라인 상세도')

# 5D Scoring
add_heading('1.4.1 5차원 할루시네이션 신뢰도 스코어링', level=3)
add_table(
    ['차원', '가중치', '산출 방법'],
    [
        ['사실 정확도', '35%', '지지 주장 수 / 전체 주장 수'],
        ['수치 정확도', '15%', '수치 지지 건수 / 수치 전체 건수'],
        ['증거 신뢰도', '20%', '신뢰도 가중 평균 (중립 패널티 적용)'],
        ['일관성', '15%', '1.0 - (모순비율 × 0.5)'],
        ['불확실성 보정', '15%', '(확정 + 정보중립×0.5) / 전체'],
    ]
)

add_figure('fig06_5d_scoring.png', '[도 6] 5차원 할루시네이션 신뢰도 스코어링 모델')

# Compliance
add_heading('1.5 규제 준수 평가 모듈 (150)', level=2)
add_para('활동 로그 관리부(151)는 AI 동작 이력을 구조화된 로그로 기록한다. 프레임워크 규칙 엔진(152)은 4개 규제 프레임워크에 대한 75개 이상의 선언적 평가 규칙을 관리한다.')

add_table(
    ['프레임워크', '규격', '주요 카테고리'],
    [
        ['EU AI Act', 'Regulation 2024/1689', '투명성, 위험관리, 데이터 거버넌스, 인간 감독, 기술적 강건성'],
        ['한국 AI 기본법', '2026년 시행', '안전성, 투명성, 공정성, 책임성, 프라이버시'],
        ['NIST AI RMF', 'AI RMF 1.0', 'Govern, Map, Measure, Manage'],
        ['OWASP LLM Top 10', 'v1.1 (2024)', '프롬프트 인젝션, DoS, 민감정보 노출 등 10개'],
    ]
)

add_figure('fig07_compliance_system.png', '[도 7] 다중 규제 프레임워크 컴플라이언스 평가 시스템')

# Integrated Control
add_heading('1.6 통합 제어 모듈 (160)', level=2)
add_para('통합 제어 모듈(160)은 사용자의 AI 콘텐츠 생성 요청 시, S1(입력 보안 검증) → S2(AI 생성) → S3(출력 보안 검증) → S4(할루시네이션 검증) → S5(인간 감독) → S6(문서 반영) → S7(컴플라이언스 로깅)의 파이프라인을 자동 실행한다.')

add_figure('fig09_sequence_diagram.png', '[도 9] 통합 제어 모듈 오케스트레이션 시퀀스 다이어그램')

# Hybrid
add_heading('1.7 온라인-오프라인 하이브리드 아키텍처', level=2)
add_para('온라인 모드: 전체 4계층 검증 파이프라인이 서버 측에서 실행된다.')
add_para('오프라인 모드: 서버 접속 불가 시, 클라이언트 측 JavaScript 엔진이 자동 활성화되어 가드레일 규칙 검증 및 수치 교차검증을 수행한다.')
add_para('서버 상태를 주기적으로 감시하며, 복구 시 자동 온라인 전환 및 미검증 항목 일괄 재검증이 가능하다.')

add_figure('fig08_hybrid_mode.png', '[도 8] 온라인-오프라인 하이브리드 운용 모드 전환 흐름도')

# Structure analysis
add_heading('1.8 문서 구조 분석 기반 프롬프트 구성', level=2)
add_para('문서 구조 분석부(121)는 S121-1(반복 패턴 탐지) → S121-2(헤더-데이터 셀 구분) → S121-3(테이블 교차참조 분석) → S121-4(문서 유형 판별)의 단계를 수행하여, 문서 구조를 보존하면서 콘텐츠만 변경하도록 제약된 프롬프트를 구성한다.')

add_figure('fig10_structure_analysis.png', '[도 10] 문서 구조 분석 기반 문맥 인식 프롬프트 구성 흐름도')

# ============================================================
# 특허청구범위
# ============================================================
doc.add_page_break()
add_heading('[특허청구범위]', level=1)

add_heading('【독립항】', level=2)

add_para('【청구항 1】', bold=True, size=11)
add_para('AI 기반 보안 검증 및 규제 준수 기능이 통합된 지능형 문서 편집 시스템에 있어서,')
add_para('문서 파일을 파싱하여 내부 문서 객체 모델로 변환하고 사용자의 편집 입력에 따라 상기 문서 객체 모델을 갱신하는 문서 파싱 및 편집 모듈(110);')
add_para('상기 문서 객체 모델의 구조를 분석하여 문서 유형을 자동 판별하고, 헤더-콘텐츠 쌍을 추출하여 문맥 인식 프롬프트를 구성한 후 외부 대규모 언어 모델(LLM) API를 호출하여 콘텐츠를 생성하는 AI 콘텐츠 생성 모듈(120);')
add_para('상기 AI 콘텐츠 생성 모듈(120)의 입력 및 출력에 대하여 복수의 방어 계층을 순차적으로 적용하여 보안 위협을 탐지 및 차단하되, 대상 언어의 언어학적 특성을 악용한 우회 공격을 탐지하는 복수의 언어 특화 탐지 서브모듈을 포함하는 다계층 보안 방어 모듈(130);')
add_para('상기 AI 콘텐츠 생성 모듈(120)이 생성한 콘텐츠에서 개별 주장을 추출하고, 각 주장에 대하여 규칙 기반 검증, 수치 교차검증, 자연어 추론 모델 기반 의미 검증을 포함하는 복수 계층의 검증 파이프라인을 적용하여 할루시네이션을 탐지하는 할루시네이션 검증 모듈(140);')
add_para('상기 AI 콘텐츠 생성 모듈(120), 보안 방어 모듈(130), 할루시네이션 검증 모듈(140)의 동작 이력을 활동 로그로 기록하고, 복수의 규제 프레임워크에 대한 평가 규칙을 적용하여 준수 여부를 자동 판정하며 증거 기반 컴플라이언스 리포트를 생성하는 규제 준수 평가 모듈(150); 및')
add_para('사용자의 AI 콘텐츠 생성 요청 시, 입력 보안 검증, AI 콘텐츠 생성, 출력 보안 검증, 할루시네이션 검증, 인간 감독, 문서 반영, 컴플라이언스 로깅의 파이프라인을 자동 실행하는 통합 제어 모듈(160);')
add_para('을 포함하는 것을 특징으로 하는, AI 기반 보안 검증 및 규제 준수 기능이 통합된 지능형 문서 편집 시스템.', bold=True)

add_para('')
add_para('【청구항 2】', bold=True, size=11)
add_para('AI 기반 보안 검증 및 규제 준수 기능이 통합된 지능형 문서 편집 방법에 있어서,')
add_para('(a) 문서 파일을 파싱하여 내부 문서 객체 모델로 변환하는 단계;')
add_para('(b) 사용자로부터 AI 콘텐츠 생성 요청 및 프롬프트를 수신하는 단계;')
add_para('(c) 상기 프롬프트에 대하여 다계층 보안 방어 모듈의 복수의 방어 계층을 순차적으로 적용하여 입력 측 보안 검증을 수행하는 단계;')
add_para('(d) 상기 문서 객체 모델의 구조를 분석하여 문서 유형을 자동 판별하고, 헤더-콘텐츠 쌍을 추출하여 문맥 인식 프롬프트를 구성한 후 외부 LLM API를 호출하여 콘텐츠를 생성하는 단계;')
add_para('(e) 생성된 콘텐츠에 대하여 상기 다계층 보안 방어 모듈을 적용하여 출력 측 보안 검증을 수행하는 단계;')
add_para('(f) 보안 검증을 통과한 콘텐츠에서 개별 주장을 추출하고, 각 주장에 대하여 규칙 기반 검증, 수치 교차검증, 자연어 추론 모델 기반 의미 검증을 포함하는 복수 계층의 검증 파이프라인을 적용하여 할루시네이션을 탐지하는 단계;')
add_para('(g) 검증 결과를 사용자에게 제시하여 인간 감독에 의한 검토, 승인, 또는 수정 입력을 수신하는 단계;')
add_para('(h) 승인된 콘텐츠를 상기 문서 객체 모델에 반영하는 단계; 및')
add_para('(i) 상기 (b) 내지 (h) 단계의 전체 과정을 활동 로그로 기록하고, 복수의 규제 프레임워크에 대한 준수 여부를 자동 평가하는 단계;')
add_para('를 포함하는 것을 특징으로 하는, AI 기반 보안 검증 및 규제 준수 기능이 통합된 지능형 문서 편집 방법.', bold=True)

add_heading('【종속항】', level=2)

claims_dep = [
    ('3', '제1항에 있어서, 상기 다계층 보안 방어 모듈(130)은, 제0계층 신뢰 경계 검증부(131); 제1계층 의도 분석부(132); 제2계층 RA 가드(133); 제3계층 ClassRAG(134); 제4계층 회로 차단기(135); 및 제5계층 행동 분석부(136);를 포함하는 6계층 방어 아키텍처로 구성되는 것을 특징으로 하는 시스템.'),
    ('4', '제3항에 있어서, 상기 행동 분석부(136)는, 일반→전문→구체 에스컬레이션의 제1 크레센도 패턴; 교육→심화→유해의 제2 크레센도 패턴; 및 신뢰→역할→제약해제의 제3 크레센도 패턴; 중 적어도 하나를 탐지하되, 세션 기반 위험도 누적과 감쇠 함수를 적용하는 것을 특징으로 하는 시스템.'),
    ('5', '제1항에 있어서, 상기 다계층 보안 방어 모듈(130)은, 키보드 매핑 탐지기; 고어 탐지기; 동형 문자 탐지기; 자모 정규화기; 은어 탐지기; 음절 역순 탐지기; 음운 변이 탐지기; 및 초성 디코더; 중 적어도 하나를 포함하는 한국어 특화 탐지 서브모듈을 포함하는 것을 특징으로 하는 시스템.'),
    ('6', '제1항에 있어서, 상기 다계층 보안 방어 모듈(130)은, 기지의 위험 패턴에 대하여 확률적 판단을 배제하고 100% 신뢰도로 차단하는 결정론적 방어 관리자(θ=0 보장)를 각 방어 계층에 포함하는 것을 특징으로 하는 시스템.'),
    ('7', '제1항에 있어서, 상기 할루시네이션 검증 모듈(140)의 수치 교차검증은, 수치와 문맥 키워드를 추출하는 단계; 문맥 키워드 기반 카테고리 할당 단계; 및 동일 카테고리 내 수치만 비교하여 유형별 허용 오차를 적용하는 단계;를 포함하는 문맥 인식 수치 교차검증인 것을 특징으로 하는 시스템.'),
    ('8', '제1항에 있어서, 상기 할루시네이션 검증 모듈(140)은, 사실 정확도(35%), 수치 정확도(15%), 증거 신뢰도(20%), 일관성(15%), 불확실성 보정(15%)의 5차원 가중 합산 스코어링을 수행하되, 중립 판정을 정보 중립과 비정보 중립으로 구분하여 차등 반영하는 것을 특징으로 하는 시스템.'),
    ('9', '제1항에 있어서, 상기 규제 준수 평가 모듈(150)은, EU AI Act, 한국 AI 기본법, NIST AI RMF, OWASP LLM Top 10 중 적어도 2개 이상에 대한 선언적 평가 규칙을 포함하고, 단일 활동 로그로부터 복수 프레임워크를 병렬 평가하는 것을 특징으로 하는 시스템.'),
    ('10', '제1항에 있어서, 상기 컴플라이언스 리포트는, 각 항목의 pass/warn/fail/n/a 판정; 활동 로그로부터 추출된 증거; 교정 지침; 및 AI 사용 요약;을 포함하는 것을 특징으로 하는 시스템.'),
    ('11', '제2항에 있어서, 상기 (f) 단계에서 할루시네이션이 탐지된 주장에 대하여, 원본 증거 기반으로 LLM을 통해 교정 텍스트를 자동 생성하여 제시하는 단계를 더 포함하는 것을 특징으로 하는 방법.'),
    ('12', '제2항에 있어서, 상기 (d) 단계에서, 반복 패턴 탐지, 헤더-데이터 셀 구분, 테이블 교차참조 분석, 문서 유형별 생성 전략 결정 및 구조 보존 제약 프롬프트 구성을 포함하는 것을 특징으로 하는 방법.'),
    ('13', '제1항에 있어서, 서버 가용 시 전체 파이프라인의 온라인 모드와, 서버 접속 불가 시 클라이언트 측 규칙 기반 및 수치 검증의 오프라인 모드를 포함하는 하이브리드 아키텍처로서, 서버 상태 감시에 따라 자동 모드 전환하는 것을 특징으로 하는 시스템.'),
    ('14', '제1항에 있어서, 상기 보안 방어 모듈(130)은, 8종 PII 탐지 스캐너; 형식 보존 가명처리 PII 프록시 엔진; 및 구성 가능 DLP 게이트웨이;를 더 포함하는 것을 특징으로 하는 시스템.'),
    ('15', '제1항에 있어서, 상기 문서 파싱 모듈(110)은, HWPX XML 파서 및 HWP 바이너리 OLE 파서를 포함하여 한글 문서 네이티브 파싱을 지원하는 것을 특징으로 하는 시스템.'),
]

for num, text in claims_dep:
    add_para(f'【청구항 {num}】', bold=True, size=11)
    add_para(text)
    add_para('')

# ============================================================
# 요약서
# ============================================================
doc.add_page_break()
add_heading('[요약서]', level=1)
add_para('【요약】', bold=True, size=11)
add_para('본 발명은 대규모 언어 모델(LLM) 기반 콘텐츠 자동 생성, 6계층 다계층 보안 방어(PALADIN 아키텍처), 4계층 할루시네이션 검증 파이프라인, 및 4대 국제 규제 프레임워크 자동 준수 평가가 단일 문서 편집 환경에 유기적으로 통합된 지능형 문서 편집 시스템 및 그 방법에 관한 것이다. 상기 시스템은 한국어를 포함한 CJK 언어 특화 우회 공격 방어, 문맥 인식 수치 교차검증, 5차원 할루시네이션 신뢰도 스코어링, 증거 기반 컴플라이언스 리포트 자동 생성, 온라인-오프라인 하이브리드 운용을 특징으로 하며, 이를 통해 AI 문서 편집의 보안성, 정확성, 규제 준수성을 종합적으로 보장한다.')

add_para('【대표도】 도 1', bold=True)
add_para('')
add_para('【대표도의 주요 부호 설명】', bold=True)
symbols = [
    '100: 지능형 문서 편집 시스템',
    '110: 문서 파싱 및 편집 모듈',
    '120: AI 콘텐츠 생성 모듈',
    '130: 다계층 보안 방어 모듈 (AEGIS)',
    '140: 할루시네이션 검증 모듈 (TruthAnchor)',
    '150: 규제 준수 평가 모듈',
    '160: 통합 제어 모듈',
    '170: 상태 관리 저장소',
    '180: 사용자 인터페이스 계층',
]
for s in symbols:
    add_para(f'• {s}')

# ============================================================
# 부록: 기술 용어 정의
# ============================================================
doc.add_page_break()
add_heading('[부록: 기술 용어 정의]', level=1)
add_table(
    ['용어', '정의'],
    [
        ['LLM', '대규모 텍스트 데이터로 사전학습된 자연어 처리 AI 모델'],
        ['할루시네이션', 'LLM이 사실과 다른 정보를 생성하는 현상'],
        ['프롬프트 인젝션', 'LLM의 시스템 지시를 조작하여 의도하지 않은 동작을 유도하는 공격'],
        ['탈옥 (Jailbreak)', 'LLM의 안전 제약을 우회하여 금지된 응답을 유도하는 공격'],
        ['NLI', '두 문장 간의 논리적 관계를 판별하는 자연어 처리 과제'],
        ['RAG', '외부 문서 검색을 통해 LLM 생성의 정확도를 향상시키는 기법'],
        ['크레센도 공격', '다회차 대화에서 점진적으로 악의적 의도를 에스컬레이션하는 공격'],
        ['PII', '개인을 식별할 수 있는 정보 (주민등록번호, 전화번호 등)'],
        ['DLP', '민감 데이터의 외부 유출을 방지하는 보안 기술'],
        ['HWPX', '한글과컴퓨터의 개방형 문서 포맷 (OWPML 기반 XML 구조)'],
        ['PALADIN', '본 발명의 6계층 순차적 방어 오케스트레이션 아키텍처'],
        ['θ=0 보장', '기지 위험 패턴에 대해 100% 차단을 보장하는 결정론적 방어'],
    ]
)

# ============================================================
# Save
# ============================================================
output_path = os.path.join(PATENT_DIR, 'patent-integrated-ai-document-system.docx')
doc.save(output_path)
print(f"Word document saved: {output_path}")
